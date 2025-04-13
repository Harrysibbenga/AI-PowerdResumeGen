from typing import Dict, Any
import os
import tempfile
from pathlib import Path
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings
import json

async def export_to_pdf(resume_data: Dict[str, Any], content: Dict[str, Any], export_path: str) -> str:
    """
    Export resume to PDF format
    
    Args:
        resume_data: Resume data from database
        content: Resume content with sections
        export_path: Path to save the PDF
        
    Returns:
        Path to the exported PDF file
    """
    try:
        # Create the HTML for the resume
        html_content = _generate_resume_html(resume_data, content)
        
        # Create directory for exports if it doesn't exist
        os.makedirs(os.path.dirname(export_path), exist_ok=True)
        
        # Create PDF from HTML
        css = CSS(string='''
            body { 
                font-family: Arial, sans-serif;
                margin: 0;
                padding: 0;
                color: #333;
            }
            .container {
                width: 8.5in;
                padding: 0.5in;
            }
            h1 { 
                text-align: center;
                font-size: 24pt;
                margin-bottom: 0.2in;
            }
            .contact-info {
                text-align: center;
                font-size: 11pt;
                margin-bottom: 0.3in;
                color: #555;
            }
            h2 {
                font-size: 14pt;
                border-bottom: 1pt solid #3498db;
                padding-bottom: 5pt;
                margin-top: 0.3in;
                color: #3498db;
            }
            .section {
                margin-bottom: 0.2in;
            }
            .job-header, .education-header {
                display: flex;
                justify-content: space-between;
                margin-bottom: 0.1in;
            }
            .job-title, .degree {
                font-weight: bold;
                font-size: 12pt;
            }
            .company, .institution {
                font-size: 11pt;
                color: #555;
            }
            .date {
                font-size: 10pt;
                color: #777;
            }
            ul {
                margin-top: 0.1in;
                padding-left: 0.2in;
            }
            li {
                margin-bottom: 0.1in;
            }
            .skills-container {
                display: flex;
                flex-wrap: wrap;
                gap: 0.1in;
            }
            .skill {
                background-color: #f0f0f0;
                padding: 0.05in 0.1in;
                border-radius: 0.05in;
                font-size: 10pt;
            }
        ''')
        
        # Generate PDF
        html = HTML(string=html_content)
        html.write_pdf(export_path, stylesheets=[css])
        
        return export_path
    
    except Exception as e:
        raise Exception(f"Error exporting resume to PDF: {str(e)}")

async def export_to_docx(resume_data: Dict[str, Any], content: Dict[str, Any], export_path: str) -> str:
    """
    Export resume to DOCX format
    
    Args:
        resume_data: Resume data from database
        content: Resume content with sections
        export_path: Path to save the DOCX file
        
    Returns:
        Path to the exported DOCX file
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Set up document properties
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        
        # Create directory for exports if it doesn't exist
        os.makedirs(os.path.dirname(export_path), exist_ok=True)
        
        # Header with name
        name = resume_data.get("profile_data", {}).get("name", "")
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run(name)
        header_run.bold = True
        header_run.font.size = Pt(18)
        
        # Contact info
        contact_info = []
        email = resume_data.get("profile_data", {}).get("email")
        if email:
            contact_info.append(email)
        
        phone = resume_data.get("profile_data", {}).get("phone")
        if phone:
            contact_info.append(phone)
        
        linkedin = resume_data.get("profile_data", {}).get("linkedin")
        if linkedin:
            contact_info.append(linkedin)
        
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(" | ".join(contact_info))
        
        # Add a line after contact info
        doc.add_paragraph("_" * 50)
        
        # Summary Section
        if content.get("summary"):
            doc.add_heading("Professional Summary", level=2)
            doc.add_paragraph(content["summary"])
        
        # Experience Section
        if resume_data.get("profile_data", {}).get("experience"):
            doc.add_heading("Experience", level=2)
            for i, job in enumerate(resume_data.get("profile_data", {}).get("experience", [])):
                p = doc.add_paragraph()
                p.add_run(f"{job.get('title', '')}").bold = True
                p.add_run(f", {job.get('company', '')}")
                
                # Date range
                date_text = f"{job.get('startDate', '')} - "
                date_text += "Present" if job.get("current") else job.get("endDate", "")
                date_paragraph = doc.add_paragraph()
                date_paragraph.add_run(date_text).italic = True
                
                # Job description bullets
                if content.get("experience") and i < len(content.get("experience", [])):
                    job_content = content["experience"][i]
                    for bullet in _split_bullet_points(job_content):
                        bullet_paragraph = doc.add_paragraph(bullet, style='List Bullet')
            
        # Skills Section
        if content.get("skills"):
            doc.add_heading("Skills", level=2)
            skills_text = ", ".join(content.get("skills", []))
            doc.add_paragraph(skills_text)
        
        # Education Section
        if resume_data.get("profile_data", {}).get("education"):
            doc.add_heading("Education", level=2)
            for edu in resume_data.get("profile_data", {}).get("education", []):
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')}").bold = True
                p.add_run(f", {edu.get('institution', '')}")
                
                # Date range
                year_text = f"{edu.get('startYear', '')} - {edu.get('endYear', '')}"
                year_paragraph = doc.add_paragraph()
                year_paragraph.add_run(year_text).italic = True
        
        # Save the document
        doc.save(export_path)
        
        return export_path
    
    except Exception as e:
        raise Exception(f"Error exporting resume to DOCX: {str(e)}")

def _generate_resume_html(resume_data: Dict[str, Any], content: Dict[str, Any]) -> str:
    """
    Generate HTML for resume
    
    Args:
        resume_data: Resume data from database
        content: Resume content with sections
        
    Returns:
        HTML string for resume
    """
    profile = resume_data.get("profile_data", {})
    
    # Contact information
    contact_parts = []
    if profile.get("email"):
        contact_parts.append(profile["email"])
    if profile.get("phone"):
        contact_parts.append(profile["phone"])
    if profile.get("linkedin"):
        contact_parts.append(profile["linkedin"])
    
    contact_info = " | ".join(contact_parts)
    
    # Experience section
    experience_html = ""
    for i, job in enumerate(profile.get("experience", [])):
        job_title = job.get("title", "")
        company = job.get("company", "")
        start_date = job.get("startDate", "")
        end_date = "Present" if job.get("current") else job.get("endDate", "")
        date_range = f"{start_date} - {end_date}"
        
        experience_html += f"""
        <div class="job">
            <div class="job-header">
                <div>
                    <div class="job-title">{job_title}</div>
                    <div class="company">{company}</div>
                </div>
                <div class="date">{date_range}</div>
            </div>
        """
        
        # Job bullet points
        if content.get("experience") and i < len(content.get("experience", [])):
            experience_html += "<ul>"
            job_content = content["experience"][i]
            for bullet in _split_bullet_points(job_content):
                experience_html += f"<li>{bullet}</li>"
            experience_html += "</ul>"
        
        experience_html += "</div>"
    
    # Skills section
    skills_html = '<div class="skills-container">'
    for skill in content.get("skills", []):
        skills_html += f'<span class="skill">{skill}</span>'
    skills_html += '</div>'
    
    # Education section
    education_html = ""
    for edu in profile.get("education", []):
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year_range = f"{edu.get('startYear', '')} - {edu.get('endYear', '')}"
        
        education_html += f"""
        <div class="education-item">
            <div class="education-header">
                <div>
                    <div class="degree">{degree}</div>
                    <div class="institution">{institution}</div>
                </div>
                <div class="date">{year_range}</div>
            </div>
        </div>
        """
    
    # Complete HTML
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Resume - {profile.get('name', '')}</title>
    </head>
    <body>
        <div class="container">
            <h1>{profile.get('name', '')}</h1>
            <div class="contact-info">{contact_info}</div>
            
            <div class="section">
                <h2>Professional Summary</h2>
                <p>{content.get('summary', '')}</p>
            </div>
            
            <div class="section">
                <h2>Experience</h2>
                {experience_html}
            </div>
            
            <div class="section">
                <h2>Skills</h2>
                {skills_html}
            </div>
            
            <div class="section">
                <h2>Education</h2>
                {education_html}
            </div>
        </div>
    </body>
    </html>
    """
    
    return html

def _split_bullet_points(text: str) -> list:
    """
    Split text into bullet points
    
    Args:
        text: Text potentially containing multiple bullet points
        
    Returns:
        List of bullet point strings
    """
    if not text:
        return []
    
    # Split by newlines or bullet points
    bullets = []
    for line in text.split('\n'):
        line = line.strip()
        if line:
            # Remove existing bullet markers
            if line.startswith('•'):
                line = line[1:].strip()
            if line.startswith('-'):
                line = line[1:].strip()
            if line.startswith('*'):
                line = line[1:].strip()
            
            bullets.append(line)
    
    return bullets